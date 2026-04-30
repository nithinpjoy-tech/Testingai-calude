"""
core/kb_ingest.py — Text extraction and chunking for the Knowledge Base.

Drop-in addition to theTest Triage Tool.
Supports: PDF, DOCX, TXT, MD, LOG, CSV, JSON
"""
from __future__ import annotations

import csv
import io
import json
import logging
from pathlib import Path
from typing import List

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {"pdf", "docx", "txt", "md", "log", "csv", "json"}


def extract_text(content: bytes, filename: str) -> str:
    """Extract plain text from uploaded file bytes."""
    ext = Path(filename).suffix.lower().lstrip(".")

    if ext == "pdf":
        return _extract_pdf(content)
    elif ext == "docx":
        return _extract_docx(content)
    elif ext == "csv":
        return _extract_csv(content)
    elif ext == "json":
        return _extract_json(content)
    elif ext in {"txt", "md", "log"}:
        return content.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported file type: .{ext}")


def _extract_pdf(content: bytes) -> str:
    try:
        import fitz  # PyMuPDF
        with fitz.open(stream=content, filetype="pdf") as doc:
            return "\n\n".join(page.get_text("text") for page in doc)
    except ImportError:
        raise ImportError("PyMuPDF not installed. Run: pip install PyMuPDF")


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")


def _extract_csv(content: bytes) -> str:
    decoded = content.decode("utf-8", errors="replace")
    rows = [", ".join(row) for row in csv.reader(io.StringIO(decoded)) if any(c.strip() for c in row)]
    return "\n".join(rows)


def _extract_json(content: bytes) -> str:
    try:
        return json.dumps(json.loads(content.decode("utf-8", errors="replace")), indent=2)
    except json.JSONDecodeError:
        return content.decode("utf-8", errors="replace")


def chunk_text(
    text: str,
    doc_id: str,
    filename: str,
    chunk_size: int = 700,
    overlap: int = 120,
) -> List[dict]:
    """
    Split text into overlapping chunks.

    Returns list of dicts:
      { id, text, doc_id, source, chunk_idx }
    """
    text = text.strip()
    if not text:
        return []

    chunks: list[dict] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk = text[start:end].strip()
        if len(chunk) > 50:
            chunks.append({
                "id": f"{doc_id}::chunk::{len(chunks)}",
                "text": chunk,
                "doc_id": doc_id,
                "source": filename,
                "chunk_idx": len(chunks),
            })
        start += chunk_size - overlap

    logger.info("'%s' → %d chunks", filename, len(chunks))
    return chunks
